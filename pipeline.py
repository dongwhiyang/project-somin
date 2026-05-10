"""
프로젝트 소민 — 5단계 에이전틱 블로그 파이프라인
1) 주제 선정 (DeepSeek/Gemma) — app.py에서 처리
2) 공공 API 수집 + DeepSeek 초안
3) Llama 3.1 70B 기술적 비판
4) DeepSeek 재수정 (결자해지)
5) LiteLLM 문체 튜닝 + SEO + 이미지 가이드
"""
import os
import re
import json
import io
import time
import base64
import requests
from litellm import completion
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv(override=True)

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
KCSC_API_KEY = os.getenv("KCSC_API_KEY", "")
LAW_API_KEY = os.getenv("LAW_API_KEY", "")
DATA_GO_KR_API_KEY = os.getenv("DATA_GO_KR_API_KEY", "")

# ─────────────────────────────────────────────
# LiteLLM 모델 매핑
# ─────────────────────────────────────────────
TUNING_MODELS = {
    "Claude 3.5 Sonnet": {
        "model": "anthropic/claude-3-5-sonnet-latest",
        "env_key": "ANTHROPIC_API_KEY",
        "icon": "🟣",
        "color": "#c084fc",
    },
    "Gemini 2.5 Flash": {
        "model": "gemini/gemini-2.5-flash",
        "env_key": "GEMINI_API_KEY",
        "icon": "⚡",
        "color": "#eab308",
    },
    "GPT-4o": {
        "model": "gpt-4o",
        "env_key": "OPENAI_API_KEY",
        "icon": "🟢",
        "color": "#4ade80",
    },
    "Grok 2": {
        "model": "xai/grok-2",
        "env_key": "XAI_API_KEY",
        "icon": "🟠",
        "color": "#fb923c",
    },
    "Llama 3.1 (NVIDIA)": {
        "model": "openai/meta/llama-3.1-70b-instruct",
        "env_key": "NVIDIA_API_KEY",
        "icon": "🟢",
        "color": "#76b900",
    },
}


def check_api_key(env_key: str) -> bool:
    val = os.getenv(env_key, "")
    return bool(val and len(val) > 5)


# ═════════════════════════════════════════════
#  정부 공공 API 연동
# ═════════════════════════════════════════════
def fetch_kcsc_data(topic: str) -> str:
    if not KCSC_API_KEY or len(KCSC_API_KEY) < 5:
        return ""
    try:
        url = f"https://kcsc.re.kr/OpenApi/CodeList?key={KCSC_API_KEY}"
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
        data = resp.json() if "json" in resp.headers.get("content-type", "") else []
        keywords = topic.replace("(", "").replace(")", "").split()[:3]
        results = []
        if isinstance(data, list):
            for item in data:
                item_str = json.dumps(item, ensure_ascii=False)
                if any(kw in item_str for kw in keywords):
                    name = item.get("name", item.get("codeName", str(item)))
                    code = item.get("code", item.get("codeNo", ""))
                    results.append(f"- [{code}] {name}")
                if len(results) >= 10:
                    break
        if results:
            return "[국가건설기준(KCSC) 관련 기준]\n" + "\n".join(results)
    except Exception:
        pass
    return ""


def fetch_law_data(topic: str) -> str:
    if not LAW_API_KEY or len(LAW_API_KEY) < 3:
        return ""
    try:
        search_terms = ["건설", "토목", "기초", "토질", "시공"]
        query = next((t for t in search_terms if t in topic), "건설기술")
        url = (
            f"https://www.law.go.kr/DRF/lawSearch.do"
            f"?OC={LAW_API_KEY}&target=law&type=JSON"
            f"&query={requests.utils.quote(query)}&display=5"
        )
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        laws = data.get("LawSearch", {}).get("law", [])
        if isinstance(laws, dict):
            laws = [laws]
        results = []
        for law in laws[:5]:
            name = law.get("법령명한글", law.get("lawName", ""))
            if name:
                results.append(f"- {name}")
        if results:
            return "[관련 법령 정보]\n" + "\n".join(results)
    except Exception:
        pass
    return ""


def fetch_qnet_data() -> str:
    if not DATA_GO_KR_API_KEY or len(DATA_GO_KR_API_KEY) < 5:
        return ""
    try:
        url = (
            f"https://apis.data.go.kr/B490007/qualExamSchd/getQualExamSchdList"
            f"?ServiceKey={DATA_GO_KR_API_KEY}&numOfRows=5&pageNo=1&dataFormat=json"
            f"&implYy=2025&qualgbCd=T"
        )
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        items = data.get("body", {}).get("items", [])
        results = []
        for item in items[:5]:
            name = item.get("jmNm", "")
            desc = item.get("implSeqNm", "")
            dt = item.get("docExamDt", "")
            if name and "토목" in name:
                results.append(f"- {name} {desc}: {dt}")
        if results:
            return "[국가기술자격 시험일정]\n" + "\n".join(results)
    except Exception:
        pass
    return ""


def collect_gov_data(topic: str) -> tuple[str, int]:
    sections, count = [], 0
    for fn, args in [
        (fetch_kcsc_data, (topic,)),
        (fetch_law_data, (topic,)),
        (fetch_qnet_data, ()),
    ]:
        r = fn(*args)
        if r:
            sections.append(r)
            count += r.count("\n")
    return "\n\n".join(sections), count


# ═════════════════════════════════════════════
#  2단계: DeepSeek V4 초안 작성
# ═════════════════════════════════════════════
def _deepseek_client():
    return OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")


def analyze_competitors_with_deepseek(topic: str, competitor_texts: list[str]) -> str:
    """스크래핑한 경쟁 블로그 본문들을 분석하여 장점과 보완점(단점)을 추출합니다."""
    if not competitor_texts:
        return ""
        
    client = _deepseek_client()
    combined_text = "\n\n---\n\n".join(competitor_texts)
    
    response = client.chat.completions.create(
        model="deepseek-v4-pro",
        messages=[
            {"role": "system", "content": """Role: 너는 최고의 SEO 및 블로그 마케팅 전문가야.
Goal: 주어진 상위 노출 경쟁 블로그 글들을 분석하여, 우리가 작성할 새로운 글이 이들을 압도할 수 있도록 벤치마킹 전략(장점 흡수, 단점 보완)을 3~4줄로 요약해 줘.
작성 규칙:
1. 경쟁 글들의 훌륭한 점(훅, 구조 등)을 파악하세요.
2. 경쟁 글들의 부족한 점(전문성, 최신 데이터 부재, 가독성 등)을 날카롭게 지적하세요.
3. 최종적으로 '새로운 글이 취해야 할 핵심 전략'을 명확히 제시하세요."""},
            {"role": "user", "content": f"주제: {topic}\n\n【경쟁 블로그 본문 샘플】\n{combined_text}\n\n위 내용들을 바탕으로 벤치마킹 분석을 작성해 주세요."},
        ],
        temperature=0.6,
        max_tokens=1500,
    )
    return response.choices[0].message.content.strip()
    
def auto_pick_topic(topics_data: dict) -> str:
    """제안된 6개의 주제 중 가장 흥미롭고 조회수가 잘 나올 법한 주제를 AI가 하나 고릅니다."""
    exam = topics_data.get("exam", [])
    field = topics_data.get("field", [])
    all_topics = exam + field
    if not all_topics:
        return ""
    
    topics_list_str = "\n".join([f"- {t}" for t in all_topics])
    
    client = _deepseek_client()
    response = client.chat.completions.create(
        model="deepseek-v4-pro",
        messages=[
            {"role": "system", "content": "너는 20년 경력의 베테랑 토목공학 기술자이자 파워 블로거야. 아래 제안된 6개의 블로그 주제 후보 중, 현재 건설 트렌드와 독자들의 관심도를 고려했을 때 가장 '조회수가 높고 유익할 것 같은' 주제를 딱 하나만 골라줘. 반드시 주제 명칭만 그대로 응답해."},
            {"role": "user", "content": f"주제 후보 목록:\n{topics_list_str}"},
        ],
        temperature=0.7,
    )
    picked = response.choices[0].message.content.strip()
    # 만약 AI가 서술형으로 답변했을 경우를 대비해 목록에 있는지 확인
    for t in all_topics:
        if t in picked:
            return t
    return all_topics[0] # 실패 시 첫 번째 주제 반환


def generate_draft(topic: str, anchor_text: str, gov_data: str = "", competitor_analysis: str = "") -> str:
    client = _deepseek_client()
    gov_section = f"\n\n【3. 정부 공공 데이터】\n{gov_data}" if gov_data else ""
    comp_section = f"\n\n【4. 경쟁 블로그 분석 및 벤치마킹 전략】\n{competitor_analysis}" if competitor_analysis else ""
    
    response = client.chat.completions.create(
        model="deepseek-v4-pro",
        messages=[
            {"role": "system", "content": """Role: 너는 스타 칼럼니스트야.
Goal: 선택된 주제와 API 데이터, 그리고 경쟁사 분석 데이터를 바탕으로 상위 노출을 싹쓸이할 전문적인 기술 칼럼 초안을 작성해 줘.

작성 규칙:
1. **자기소개 금지**: "안녕하세요, 10년 차 블로거입니다" 같은 불필요한 서두는 절대 쓰지 마세요. 바로 본론으로 들어가세요.
2. 서론, 본론(핵심 내용 3~4개 소제목), 결론 구조로 작성하세요.
3. 각 소제목은 ## 으로 시작하세요.
4. 전문 용어는 괄호로 쉬운 설명을 덧붙이세요.
5. 분량은 1500~2000자 내외로 작성하세요.
6. 기출문제 데이터를 참고하여 실제 시험 출제 내용을 자연스럽게 녹이세요.
7. 정부 공공 데이터(건설기준, 법령, 시험일정)가 제공되면 근거자료로 인용하세요.
8. 복잡한 수치/법령 조항은 HTML 표나 불렛 포인트로 정리하세요.
9. **경쟁사 분석 전략 반영**: 제공된 벤치마킹 전략을 철저히 반영하여 기존 상위 글들의 단점을 완벽히 보완하세요."""},
            {"role": "user", "content": f"블로그 주제: {topic}\n\n【1. 기출문제 데이터】\n{anchor_text[:8000]}{gov_section}{comp_section}\n\n위 데이터를 종합 참고하여 블로그 포스팅 초안을 작성해 주세요."},
        ],
        temperature=0.7,
        max_tokens=4000,
    )
    return response.choices[0].message.content.strip()


# ═════════════════════════════════════════════
#  3단계: Llama 3.1 70B 기술적 비판
# ═════════════════════════════════════════════
def critique_with_qwen(draft: str, topic: str) -> str:
    """NVIDIA Llama 3.1 70B Instruct가 편집장으로서 초안을 기술적으로 비판합니다."""
    kwargs = {
        "model": "openai/meta/llama-3.1-70b-instruct",
        "api_base": "https://integrate.api.nvidia.com/v1",
        "api_key": os.getenv("NVIDIA_API_KEY", ""),
        "messages": [
            {"role": "system", "content": """Role: 너는 10년 차 전문 블로거이자 대학교수, 국토부 전문 자문위원이야.
Goal: 
-	초안의 기술적 논리와 가독성을 매섭게 비판하고 입력된 초안이 '실제 현장 경험이 없는 초보가 쓴 얕은 글'인지 예리하게 감별해. 최고 수준의 기술사 답안지나 현직 엔지니어의 실무 칼럼 수준으로 초안을 끌어올리기 위해, 다음 단계의 AI(수정 모델)가 구체적으로 무엇을 고치고 추가해야 할지 가혹하고 전문적인 수정 지시를 내려.
-	 그리고 글의 이해를 돕기 위해 시각 자료가 필요한 곳을 1~3군데 선정해. 각 위치마다 [IMAGE_PROMPT: 여기에 영어로 이미지 생성용 프롬프트 작성] 형태로 정확히 어떤 사진/그림이 들어가야 할지 아이디어를 제시해 줘. 너무 많은 사진은 지양해.

프롬프트(IMAGE_PROMPT) 작성 가이드:
1. **무조건 실사 사진(Photographic)만 요청**: 건설 현장의 전경, 작업 중인 기계, 자연스러운 풍경 등 "글자가 전혀 필요 없는 실제 사진"만 묘사하세요.
2. **모든 문자/언어 배제**: 영어, 한글, 중국어를 포함한 그 어떤 형태의 글자도 이미지에 나타나지 않도록 프롬프트를 구성하세요.
3. **도표/도면/인포그래픽 금지**: 글자가 조금이라도 들어갈 여지가 있는 도표(Chart), 설계도(Blueprint), 표지판, 전광판 등은 완전히 제외하세요.
4. **품질 키워드 필수**: 'High quality photograph', 'photorealistic', 'cinematic lighting', 'no text', 'no letters'를 프롬프트 끝에 반드시 포함하세요.

## 전문 비판 및 검토 기준 (필수)
1. 기술적 오류 점검: 공학적 개념, 수치, 공법 설명의 정확성 검증 (KCS 등 건설기준 준수 여부).
2. 현장감 부재 지적: 단순 이론만 나열되었는가? 시공 시 주의사항, 품질/안전/환경 관리 대책, 시공 간섭(인터페이스) 등 '현장 실무자'의 관점이 누락되었는지 짚어내고 추가를 지시해.
3. 비교 분석 요구(엔지니어링 마인드): 단일 공법 설명에 그치지 말고, 반드시 타 공법과의 '비교표(시공성, 경제성, 유지관리성 등)'를 추가하도록 지시해. 
4. 데이터 기반: 가능하면 관련 법령, 표준품셈, 최근 사고 사례 등을 인용하여 팩트체크를 강화하고: KCS(국가건설기준), 법규, 강도, 배합비, 안전율 등 공학적 수치나 공식에 오류가 없는지 팩트 체크해. 
5. 구조적 문제: 기술사 답안지의 정석인 [개요 -> 원리/메커니즘 -> 특징 및 비교 -> 현장 시공 시 유의사항(결론)] 구조를 갖추도록 목차 재구성을 지시해.
6. **독자 관점**: 이해하기 어려운 표현, 불필요한 전문어 남용

7. 추가적인 문체 튜닝 지시
- 스토리텔링 구조: 도입부의 훅(Hook)이 강력한지, 단락 간 연결이 자연스러운지 점검.
-  가독성 최적화: 너무 긴 문장은 나누고, 복잡한 나열은 리스트나 표로 변환하도록 유도.
- 전문가의 품격: 말투는 친절하되(경어체), 기술사로서의 신뢰감이 느껴지는 단어 선택(예: '함께 살펴보겠습니다' 등)을 권장.

반드시 아래 형식으로 응답하세요:

## 🏗️ 채점위원 총평 (이 글의 현재 수준과 치명적인 약점이 무엇인지 2~3줄로 냉정하게 평가)

## ❌ 기술적 오류 및 논리적 허점
1. [지적 사항] -> [이렇게 수정하라고 구체적인 공학적 지침 제시]
2. [지적 사항] -> [수정 지침]

## 👷 현장 실무 및 심화 내용 추가 지시 (필수 반영)
- [어떤 시공상 유의사항이나 현장 사례를 추가할 것인지 지시]
  - [어떤 항목들의 비교 분석표를 추가할 것인지 지시]

## 🖼️ 추천 시각 자료 (1~3개)
1. 문맥 설명: ... -> [IMAGE_PROMPT: Detailed English Prompt Here]"""},
            {"role": "user", "content": f"블로그 주제: {topic}\n\n아래 초안을 비판해 주세요:\n\n{draft}"},
        ],
        "temperature": 0.4,
        "max_tokens": 3000,
    }
    response = completion(**kwargs)
    return response.choices[0].message.content.strip()


# ═════════════════════════════════════════════
#  3단계: DeepSeek 재수정 (통합 비판 수용)
# ═════════════════════════════════════════════
def revise_with_deepseek(draft: str, critique: str, topic: str, image_paths: list = None) -> str:
    """통합 비판을 반영하여 DeepSeek가 블로그 형식의 최종본을 작성합니다."""
    client = _deepseek_client()
    
    img_info = ""
    if image_paths:
        img_info = "\n\n【사용 가능한 삽화 목록】\n"
        for i, path in enumerate(image_paths):
            img_info += f"- 이미지 {i}: {path} 관련 삽화 -> 본문에 [IMAGE_PLACEHOLDER_{i}] 라고 기재하세요.\n"
        img_info += "\n(중요: 위 [IMAGE_PLACEHOLDER_idx] 태그를 본문의 가장 적절한 위치에 삽입하세요. 해당 위치에 나중에 실제 이미지가 들어갈 예정입니다.)"
    else:
        img_info = "\n\n(참고: 현재 사용 가능한 이미지 파일이 없습니다.)"

    response = client.chat.completions.create(
        model="deepseek-v4-pro",
        messages=[
            {"role": "system", "content": """Role: 너는 스타 IT/기술 블로거이자 전문 에디터야.
Goal: 제공된 초안과 비판 내용을 바탕으로 완성형 블로그 포스팅을 작성해 줘.

작성 및 수정 규칙:
1. **자기소개 금지**: 글 시작 부분에 본인 소개 문구는 절대 넣지 마세요.
2. **문체**: 친근한 경어체(~해요, ~입니다)를 사용하세요. 
3. **이미지 배치**: 제공된 [IMAGE_PLACEHOLDER_idx] 태그를 본문의 가장 맥락에 맞는 위치(소제목 아래 등)에 자연스럽게 삽입하세요.
4. **구조**: ## 소제목, 표(table), 리스트를 활용하여 가독성을 높이세요.
5. **최종본만 출력**: 다른 설명 없이 본문만 응답하세요."""},
            {"role": "user", "content": f"주제: {topic}\n\n【원본 초안】\n{draft}\n\n【통합 비판 가이드】\n{critique}{img_info}\n\n위 내용을 바탕으로 블로그 글을 작성해 주세요."},
        ],
        temperature=0.6,
        max_tokens=4000,
    )
    final_text = response.choices[0].message.content.strip()
    
    # 이미지 호스팅 업로드 헬퍼
    def upload_to_freeimage(img_path):
        try:
            with open(img_path, "rb") as f:
                b64_data = base64.b64encode(f.read()).decode('utf-8')
            res = requests.post(
                "https://freeimage.host/api/1/upload", 
                data={'key': '6d207e02198a847aa98d0a2a901485a5', 'action': 'upload', 'source': b64_data},
                timeout=30
            )
            if res.status_code == 200:
                return res.json().get("image", {}).get("url", "")
        except Exception as e:
            print(f"Image upload failed: {e}")
        return ""

    # 사후 처리: 플레이스홀더를 실제 이미지 태그로 치환
    if image_paths:
        for i, path in enumerate(image_paths):
            try:
                public_url = upload_to_freeimage(path)
                if public_url:
                    img_tag = f"![삽화 {i}]({public_url})"
                else:
                    # 실패 시 기존 Base64 폴백
                    with open(path, "rb") as f:
                        b64 = base64.b64encode(f.read()).decode('utf-8')
                        img_tag = f"![삽화 {i}](data:image/png;base64,{b64})"
                        
                final_text = final_text.replace(f"[IMAGE_PLACE_HOLDER_{i}]", img_tag)
                final_text = final_text.replace(f"[IMAGE_PLACEHOLDER_{i}]", img_tag)
            except Exception:
                continue
    
    # 기타 잔여 태그 제거
    import re
    final_text = re.sub(r'\[IMAGE_PROMPT:.*?\]', '', final_text).strip()
    
    return final_text


# ═════════════════════════════════════════════
#  이미지 생성: Gemini 2.5 Flash Image (함수명 generate_image_from_nvidia 는 호환 유지)
# ═════════════════════════════════════════════
def generate_image_from_nvidia(prompt: str, idx: int) -> str:
    """Google Gemini API `gemini-2.5-flash-image` 로 삽화를 생성합니다 (GEMINI_API_KEY).
    (이름은 과거 NVIDIA Flux 호출과의 호환을 위해 유지합니다.)
    """
    strict_prompt = (
        f"{prompt}, strictly no text, no letters, no words, no symbols, "
        "no signs, no labels, photorealistic, clean composition"
    )
    api_key = (os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY") or "").strip()
    if not api_key:
        print("Gemini image: GEMINI_API_KEY 또는 GOOGLE_API_KEY 없음")
        return ""

    url = (
        "https://generativelanguage.googleapis.com/v1beta/models/"
        "gemini-2.5-flash-image:generateContent"
    )
    payload = {
        "contents": [{"parts": [{"text": strict_prompt}]}],
        "generationConfig": {"responseModalities": ["TEXT", "IMAGE"]},
    }
    try:
        response = requests.post(url, params={"key": api_key}, json=payload, timeout=120)
        if response.status_code != 200:
            print(f"Gemini image HTTP {response.status_code}: {response.text[:400]}")
            return ""
        data = response.json()
        candidates = data.get("candidates") or []
        if not candidates:
            print(f"Gemini image: candidates 없음 — {str(data)[:400]}")
            return ""
        parts = (candidates[0].get("content") or {}).get("parts") or []
        for part in parts:
            inline = part.get("inlineData") or part.get("inline_data")
            if not inline:
                continue
            raw_b64 = inline.get("data") or ""
            mime = (inline.get("mimeType") or inline.get("mime_type") or "image/png").lower()
            ext = ".jpg" if "jpeg" in mime or "jpg" in mime else ".png"
            os.makedirs("images", exist_ok=True)
            filename = f"images/generated_{int(time.time())}_{idx}{ext}"
            with open(filename, "wb") as f:
                f.write(base64.b64decode(raw_b64))
            print(f"Successfully saved image (Gemini): {filename}")
            return filename
        print("Gemini image: 응답에 이미지 바이너리 없음")
        return ""
    except Exception as e:
        print(f"Gemini image generation failed: {e}")
        return ""

# ═════════════════════════════════════════════
#  4단계: Mistral AI 문체 튜닝
# ═════════════════════════════════════════════
def tune_with_mistral(draft: str, topic: str) -> str:
    """NVIDIA Mistral AI를 활용하여 최종 문체 튜닝을 수행합니다."""
    kwargs = {
        "model": "openai/mistralai/mistral-large-2407", # Use mistral-large-2407 mapped in NVIDIA API
        "api_base": "https://integrate.api.nvidia.com/v1",
        "api_key": os.getenv("NVIDIA_API_KEY", ""),
        "messages": [
            {"role": "system", "content": "Role: 너는 20년 경력의 베테랑 토목공학 기술자이자 파워 블로거야. Task: 아래의 초안을 바탕으로, 전문 용어의 정확성은 유지하되 독자들이 읽기 편하도록 '경어체(~해요, ~입니다)'를 섞어 매끄럽게 다듬어줘. Constraint: 문장은 너무 길지 않게 끊어서 작성하고, 기술적인 신뢰감이 느껴지는 문체를 사용해줘."},
            {"role": "user", "content": f"블로그 주제: {topic}\n\n아래 초안을 매끄러운 경어체로 다듬어줘:\n\n{draft}"},
        ],
        "temperature": 0.4,
        "max_tokens": 4000,
    }
    try:
        response = completion(**kwargs)
        return response.choices[0].message.content.strip()
    except Exception as e:
        # Fallback to mixtral if mistral-large-2407 is not available
        kwargs["model"] = "openai/mistralai/mixtral-8x22b-instruct-v0.1"
        response = completion(**kwargs)
        return response.choices[0].message.content.strip()



# ═════════════════════════════════════════════
#  5단계 부가: SEO 메타데이터 생성
# ═════════════════════════════════════════════
def generate_seo_metadata(topic: str, final_text: str) -> dict:
    """SEO 태그 5개, 메타 설명, 이미지 Alt Text 3개를 생성합니다."""
    client = _deepseek_client()
    response = client.chat.completions.create(
        model="deepseek-v4-pro",
        messages=[
            {"role": "system", "content": """블로그 SEO 전문가입니다. 반드시 아래 JSON 형식으로만 응답하세요.

{
  "seo_tags": ["태그1", "태그2", "태그3", "태그4", "태그5"],
  "meta_description": "150자 내외의 메타 설명문",
  "image_alt_texts": ["이미지1 대체텍스트", "이미지2 대체텍스트", "이미지3 대체텍스트"]
}"""},
            {"role": "user", "content": f"주제: {topic}\n\n본문 요약:\n{final_text[:2000]}\n\n네이버/구글 검색 최적화 태그 5개, 메타 설명문(150자), 이미지 Alt Text 3개를 JSON으로 생성하세요."},
        ],
        temperature=0.5,
        max_tokens=800,
    )
    raw = response.choices[0].message.content.strip()
    
    # 1. 마크다운 코드 블록 및 불필요한 텍스트 제거
    raw = re.sub(r'```(?:json)?', '', raw)
    raw = re.sub(r'```', '', raw)
    
    # 2. 가장 바깥쪽 { } 추출
    try:
        start = raw.find('{')
        end = raw.rfind('}')
        if start != -1 and end != -1:
            raw = raw[start:end+1]
            
        # 3. 제어 문자 제거 및 파싱
        raw_clean = re.sub(r'[\x00-\x1F\x7F]', '', raw)
        # 따옴표 내부의 줄바꿈 처리 (JSON 위반 방지)
        raw_clean = raw_clean.replace('\n', '\\n').replace('\r', '\\r')
        # 하지만 키값 사이의 줄바꿈은 유지해야 하므로 다시 복구 (복잡하지만 안전하게)
        raw_clean = raw_clean.replace('}\\n', '}').replace('",\\n', '",')
        
        return json.loads(raw_clean)
    except Exception as e:
        print(f"JSON Parsing failed: {e}")
        return {
            "seo_tags": [topic[:10], "건설", "실무", "기술사"],
            "meta_description": f"{topic}에 대한 전문 기술 가이드입니다.",
            "image_alt_texts": ["관련 이미지"]
        }


# ═════════════════════════════════════════════
#  5단계: LiteLLM 문체 튜닝
# ═════════════════════════════════════════════
def tune_with_model(text: str, topic: str, model_key: str) -> str:
    info = TUNING_MODELS[model_key]
    kwargs = {
        "model": info["model"],
        "messages": [
            {"role": "system", "content": """전문 에디터입니다. 블로그 최종본의 문체를 윤문하세요.
1. 기술적 정확성 유지, 더 매력적이고 읽기 쉬운 문체로.
2. 구조(서론/본론/결론) 유지, 표현과 흐름 개선.
3. 독자의 관심을 끄는 도입부와 마무리.
4. ## 소제목 형식 유지.
5. HTML 표/리스트가 있으면 그대로 유지."""},
            {"role": "user", "content": f"주제: {topic}\n\n아래 최종본을 윤문하세요:\n\n{text}"},
        ],
        "temperature": 0.7,
        "max_tokens": 4000,
    }
    
    # NVIDIA API 특수 처리 (OpenAI 호환 방식)
    if info["env_key"] == "NVIDIA_API_KEY":
        kwargs["api_base"] = "https://integrate.api.nvidia.com/v1"
        kwargs["api_key"] = os.getenv("NVIDIA_API_KEY", "")
        
    response = completion(**kwargs)
    return response.choices[0].message.content.strip()


# ═════════════════════════════════════════════
#  워드 파일 생성 (5단계 전체 포함)
# ═════════════════════════════════════════════
def create_docx(topic, draft, critique, revised, tuned_results, seo_data) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    title = doc.add_heading(f"📝 {topic}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # 2단계: 초안
    doc.add_heading("━━ 2단계: 초안 (DeepSeek V4) ━━", level=1)
    doc.add_paragraph(draft)
    doc.add_page_break()

    # 3단계: 비판
    doc.add_heading("━━ 3단계: 기술적 비판 (Llama 3.1 70B) ━━", level=1)
    doc.add_paragraph(critique)
    doc.add_page_break()

    # 4단계: 최종 수정본
    doc.add_heading("━━ 4단계: 최종 수정본 (DeepSeek V4) ━━", level=1)
    doc.add_paragraph(revised)
    doc.add_page_break()

    # 5단계: 튜닝
    for name, text in tuned_results.items():
        doc.add_heading(f"━━ 5단계: 문체 튜닝 ({name}) ━━", level=1)
        doc.add_paragraph(text)
        doc.add_page_break()

    # SEO
    if seo_data:
        doc.add_heading("━━ SEO 메타데이터 ━━", level=1)
        doc.add_paragraph(f"태그: {', '.join(seo_data.get('seo_tags', []))}")
        doc.add_paragraph(f"메타 설명: {seo_data.get('meta_description', '')}")
        doc.add_paragraph("이미지 Alt Text:")
        for alt in seo_data.get("image_alt_texts", []):
            doc.add_paragraph(f"  • {alt}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def call_llama_for_topics(anchor_text: str, news_text: str) -> dict:
    """
    NVIDIA Llama 3 70B 모델로 수험 4개 + 실무 2개, 총 6개 주제 생성.
    반환: {"exam": ["주제1",...], "field": ["주제5","주제6"]}
    """
    from litellm import completion
    import os, json, re

    system_prompt = """Role: 너는 10년 차 건설 전문 파워블로거이자, 현직 토목 전문 공무원이야.
Goal: 제공된 기출문제와 뉴스 데이터를 분석해서, 실무자들이 무조건 클릭할 수밖에 없는 '실무 밀착형' 블로그 주제 6개를 제안해 줘.

규칙:
1. 주제의 구체성: "시공 관리" 같은 추상적 제목 금지. "BIM으로 깎아먹는 공기 10일 단축법"처럼 구체적이어야 함.
2. 클릭 유도 요소: 숫자를 활용(예: 3가지 팁, 5대 공법), 질문형 제목, 혹은 '현직자만 아는' 뉘앙스를 풍길 것.
3. 데이터 융합: 기출문제의 기술적 깊이와 최신 뉴스의 시의성을 섞어서 '지금 바로 필요한 정보'로 만들 것.
4. "exam" 키: 기출 기반 수험 주제 4개 (예: (토목시공기술사) 불합격자만 모르는 '교량 하부 구조' 서술 비법)
5. "field" 키: 뉴스 기반 실무 가이드 주제 2개 (예: (실무가이드) 2026 표준품셈 개정안, 이것 모르면 설계변경 다 틀립니다)
6. 반드시 아래 JSON 형식으로만 응답 (다른 텍스트 금지):

{
  "exam": [
    "(자격증명) 구체적이고 매력적인 제목 1",
    "(자격증명) 구체적이고 매력적인 제목 2",
    "(자격증명) 구체적이고 매력적인 제목 3",
    "(자격증명) 구체적이고 매력적인 제목 4"
  ],
  "field": [
    "(실무가이드) 구체적이고 매력적인 제목 5",
    "(실무가이드) 구체적이고 매력적인 제목 6"
  ]
}"""

    user_prompt = f"""【기출문제 데이터】\n{anchor_text[:6000]}\n\n【최신 실무 뉴스 헤드라인】\n{news_text[:2000]}\n\n위 데이터를 종합해서 수험 주제 4개 + 실무 주제 2개를 JSON으로 응답하세요."""

    kwargs = {
        "model": "openai/meta/llama-3.1-70b-instruct",
        "api_base": "https://integrate.api.nvidia.com/v1",
        "api_key": os.getenv("NVIDIA_API_KEY", ""),
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.8,
        "max_tokens": 1500,
    }

    try:
        response = completion(**kwargs)
        raw = response.choices[0].message.content.strip()
        
        # JSON 추출 (코드블록 제거)
        match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", raw, re.DOTALL)
        if match:
            raw = match.group(1)
        else:
            start, end = raw.find("{"), raw.rfind("}")
            if start != -1 and end != -1:
                raw = raw[start:end + 1]

        parsed = json.loads(raw.strip())
    except Exception as e:
        print(f"Error in call_llama_for_topics: {e}")
        parsed = {"exam": [], "field": []}

    # 최소 보장
    if "exam" not in parsed: parsed["exam"] = []
    if "field" not in parsed: parsed["field"] = []
    return parsed
